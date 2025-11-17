import io
import logging
from typing import Optional
import PyPDF2
import docx
import pandas as pd
from openpyxl import load_workbook

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for extracting text from various document formats"""
    
    async def extract_text(self, file_content: bytes, file_extension: str) -> str:
        """
        Extract text from document based on file type
        
        Args:
            file_content: Binary content of the file
            file_extension: File extension (e.g., '.pdf', '.docx')
        
        Returns:
            Extracted text content
        """
        try:
            if file_extension == '.pdf':
                return await self._extract_from_pdf(file_content)
            elif file_extension in ['.docx', '.doc']:
                return await self._extract_from_docx(file_content)
            elif file_extension in ['.xlsx', '.xls']:
                return await self._extract_from_excel(file_content)
            else:
                logger.warning(f"Unsupported file extension: {file_extension}")
                return ""
        
        except Exception as e:
            logger.error(f"Error extracting text from {file_extension}: {str(e)}")
            return f"Error extracting text: {str(e)}"
    
    async def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
            
            extracted_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(extracted_text)} characters from PDF")
            return extracted_text
        
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            return f"Error reading PDF: {str(e)}"
    
    async def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from Word document"""
        try:
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    table_text.append(row_text)
                
                if table_text:
                    text_parts.append("\n[TABLE]\n" + "\n".join(table_text) + "\n[/TABLE]")
            
            extracted_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(extracted_text)} characters from DOCX")
            return extracted_text
        
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            return f"Error reading Word document: {str(e)}"
    
    async def _extract_from_excel(self, file_content: bytes) -> str:
        """Extract text from Excel file"""
        try:
            excel_file = io.BytesIO(file_content)
            
            # Try with pandas first (works for both xls and xlsx)
            try:
                # Read all sheets
                excel_data = pd.read_excel(excel_file, sheet_name=None)
                
                text_parts = []
                for sheet_name, df in excel_data.items():
                    text_parts.append(f"--- Sheet: {sheet_name} ---")
                    
                    # Convert dataframe to readable text
                    # Include column headers and first 100 rows
                    df_sample = df.head(100)
                    text_parts.append(df_sample.to_string(index=False))
                    
                    # Add summary statistics for numeric columns
                    numeric_cols = df.select_dtypes(include=['number']).columns
                    if len(numeric_cols) > 0:
                        text_parts.append("\n[SUMMARY STATISTICS]")
                        text_parts.append(df[numeric_cols].describe().to_string())
                
                extracted_text = "\n\n".join(text_parts)
                logger.info(f"Extracted {len(extracted_text)} characters from Excel")
                return extracted_text
            
            except Exception as e:
                logger.error(f"Pandas Excel extraction failed: {str(e)}")
                
                # Fallback to openpyxl for xlsx
                try:
                    excel_file.seek(0)
                    wb = load_workbook(excel_file, read_only=True, data_only=True)
                    
                    text_parts = []
                    for sheet_name in wb.sheetnames:
                        sheet = wb[sheet_name]
                        text_parts.append(f"--- Sheet: {sheet_name} ---")
                        
                        # Read first 100 rows
                        for row_idx, row in enumerate(sheet.iter_rows(max_row=100, values_only=True)):
                            if any(cell is not None for cell in row):
                                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                                text_parts.append(row_text)
                            
                            if row_idx >= 100:
                                break
                    
                    extracted_text = "\n".join(text_parts)
                    logger.info(f"Extracted {len(extracted_text)} characters from Excel (openpyxl)")
                    return extracted_text
                
                except Exception as openpyxl_error:
                    logger.error(f"OpenPyXL Excel extraction also failed: {str(openpyxl_error)}")
                    return f"Error reading Excel file: {str(openpyxl_error)}"
        
        except Exception as e:
            logger.error(f"Excel extraction error: {str(e)}")
            return f"Error reading Excel file: {str(e)}"
    
    async def extract_financial_data(self, text: str) -> dict:
        """
        Extract structured financial data from text
        
        Args:
            text: Extracted text from financial documents
        
        Returns:
            Dictionary with extracted financial metrics
        """
        # This is a simplified version
        # In production, you'd use more sophisticated NLP or regex patterns
        
        financial_data = {
            "currency_mentions": [],
            "numeric_values": [],
            "keywords_found": []
        }
        
        # Common financial keywords
        keywords = [
            "revenue", "profit", "loss", "assets", "liabilities",
            "turnover", "balance", "equity", "cash flow", "budget"
        ]
        
        text_lower = text.lower()
        
        for keyword in keywords:
            if keyword in text_lower:
                financial_data["keywords_found"].append(keyword)
        
        # Extract currency mentions (USD, EUR, GBP, PKR, etc.)
        import re
        currency_pattern = r'\b(USD|EUR|GBP|PKR|INR|CNY|JPY)\b'
        currencies = re.findall(currency_pattern, text, re.IGNORECASE)
        financial_data["currency_mentions"] = list(set(currencies))
        
        # Extract numeric values with currency symbols or keywords
        # Pattern for numbers with optional commas and decimals
        number_pattern = r'[\$£€₹]\s*[\d,]+\.?\d*|\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\s*(?:million|billion|thousand|k|m|b)?'
        numbers = re.findall(number_pattern, text, re.IGNORECASE)
        financial_data["numeric_values"] = numbers[:20]  # First 20 numbers
        
        return financial_data
    
    async def extract_technical_capabilities(self, text: str) -> dict:
        """
        Extract technical capabilities and experience from text
        
        Args:
            text: Extracted text from technical documents
        
        Returns:
            Dictionary with technical information
        """
        technical_data = {
            "technologies_mentioned": [],
            "certifications": [],
            "experience_indicators": []
        }
        
        # Common technology keywords
        tech_keywords = [
            "python", "java", "javascript", "react", "angular", "vue",
            "docker", "kubernetes", "aws", "azure", "gcp",
            "microservices", "api", "database", "postgresql", "mongodb",
            "machine learning", "ai", "devops", "ci/cd"
        ]
        
        # Certification keywords
        cert_keywords = [
            "iso", "cmmi", "pci", "hipaa", "soc 2", "gdpr",
            "certified", "certification", "accreditation"
        ]
        
        text_lower = text.lower()
        
        for tech in tech_keywords:
            if tech in text_lower:
                technical_data["technologies_mentioned"].append(tech)
        
        for cert in cert_keywords:
            if cert in text_lower:
                technical_data["certifications"].append(cert)
        
        # Look for experience indicators
        import re
        experience_pattern = r'(\d+)\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp\.?)'
        experiences = re.findall(experience_pattern, text_lower)
        if experiences:
            technical_data["experience_indicators"] = [f"{exp} years" for exp in experiences]
        
        return technical_data